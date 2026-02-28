import os
from docx import Document
from fpdf import FPDF
import openpyxl

DATA_DIR = "data/raw"

def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path)

# --- TXT Documents ---
def create_txt_rules():
    content = """République Tunisienne
Ministère de l'Enseignement Supérieur et de la Recherche Scientifique
Institut Supérieur (IIT/NAU)

Règlement intérieur de l'université
Date: 01/09/2023

Article 1: Assiduité et Absences
La présence à tous les cours, TD et TP est obligatoire. Toute absence doit être justifiée dans un délai de 48 heures au service de la scolarité en présentant un certificat médical visé par le médecin de la santé publique ou un justificatif officiel.
Trois absences non justifiées consécutives dans un module entraînent l'exclusion de l'étudiant pour la session principale de ce module.

Article 2: Rattrapage
Les examens de rattrapage sont organisés à la fin de l'année universitaire pour les étudiants n'ayant pas validé la session principale. L'inscription aux examens de rattrapage est automatique. Les notes inférieures à 10/20 peuvent être repassées.

Article 3: Discipline
Toute tentative de fraude lors d'un examen entraînera automatiquement la traduction de l'étudiant devant le conseil de discipline de l'université.
"""
    with open(os.path.join(DATA_DIR, "reglement_interieur.txt"), "w", encoding="utf-8") as f:
        f.write(content)

# --- DOCX Documents ---
def create_docx_inscription():
    doc = Document()
    doc.add_heading("Procédures d'Inscription et de Réinscription", 0)
    
    doc.add_paragraph("Service de la scolarité")
    doc.add_paragraph("Année universitaire 2024-2025")
    
    doc.add_heading("Nouvelle Inscription", level=1)
    doc.add_paragraph("Pour les nouveaux bacheliers, l'inscription se fait entièrement en ligne sur le site officiel www.inscription.tn. Les étudiants doivent payer les frais d'inscription via leur carte e-dinar. Le dossier papier (comprenant 4 photos d'identité, une copie de la CIN, et le reçu de paiement) doit être déposé au service de la scolarité avant le 15 septembre.")
    
    doc.add_heading("Réinscription", level=1)
    doc.add_paragraph("La réinscription pour les étudiants en deuxième et troisième année est soumise à la validation de l'année précédente. Les frais de réinscription s'élèvent à 80 DT pour la licence et 120 DT pour le cycle ingénieur.")
    
    doc.add_heading("Attestation de scolarité", level=1)
    doc.add_paragraph("La demande d'attestation de scolarité se fait au guichet numéro 1. L'attestation est délivrée 48h après le dépôt de la demande munie de la copie de la carte d'étudiant.")
    
    doc.save(os.path.join(DATA_DIR, "procedures_inscription.docx"))

# --- PDF Documents ---
def create_pdf_calendar():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Calendrier Universitaire 2024-2025", ln=1, align='C')
    
    pdf.set_font("Arial", size=12)
    content = [
        "Département des affaires académiques",
        "",
        "Semestre 1:",
        "- Début des cours: 02 Septembre 2024",
        "- Vacances d'hiver: 16 Décembre 2024 au 01 Janvier 2025",
        "- Examens session principale: 06 Janvier 2025 au 18 Janvier 2025",
        "",
        "Semestre 2:",
        "- Début des cours: 27 Janvier 2025",
        "- Vacances de printemps: 17 Mars 2025 au 30 Mars 2025",
        "- Examens session principale: 02 Juin 2025 au 14 Juin 2025",
        "",
        "Session de Rattrapage:",
        "- Déroulement des examens: 23 Juin 2025 au 05 Juillet 2025"
    ]
    for line in content:
        pdf.cell(200, 10, txt=line, ln=1)
        
    pdf.output(os.path.join(DATA_DIR, "calendrier_universitaire.pdf"))

def create_pdf_bourses():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", 'B', 16)
    pdf.cell(200, 10, txt="Conditions d'Obtention des Bourses et Prêts", ln=1, align='C')
    
    pdf.set_font("Arial", size=12)
    content = [
        "Département des affaires estudiantines",
        "Date de publication: 10/10/2023",
        "",
        "1. Demande de Bourse Universitaire",
        "La demande de bourse s'effectue sur le site de l'Office des Oeuvres Universitaires (OOUN).",
        "Les critères d'attribution prennent en compte le revenu des parents, le nombre d'enfants",
        "à charge et l'éloignement géographique. Le dernier délai de dépôt est le 31 octobre.",
        "",
        "2. Renouvellement de la Bourse",
        "Le renouvellement n'est pas automatique. L'étudiant doit justifier sa réussite à",
        "la session principale ou de rattrapage. Les redoublants perdent leur droit à la bourse.",
        "",
        "3. Prêts Universitaires",
        "Les étudiants non boursiers peuvent postuler pour un prêt universitaire. Le dossier",
        "nécessite la signature d'un garant solvable."
    ]
    for line in content:
        pdf.cell(200, 10, txt=line, ln=1)
        
    pdf.output(os.path.join(DATA_DIR, "bourses_et_prets.pdf"))

# --- XLSX Documents ---
def create_xlsx_stages():
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet.title = "Procédures de Stage"
    
    sheet["A1"] = "Département des Stages"
    sheet["A2"] = "Date de mise à jour: 05/02/2024"
    
    sheet["A4"] = "Type de Stage"
    sheet["B4"] = "Durée Minimale"
    sheet["C4"] = "Période"
    sheet["D4"] = "Document Requis"
    
    data = [
        ("Stage Ouvrier (1ère année)", "1 mois minimum", "Été (Juillet-Août)", "Lettre d'affectation, convention signée (3 copies)"),
        ("Stage Technicien (2ème année)", "1 mois minimum", "Été (Juillet-Août)", "Lettre d'affectation, convention signée (3 copies)"),
        ("Stage de Fin d'Études (PFE)", "4 à 6 mois", "Février-Juin", "Cahier des charges validé, convention signée (3 copies)")
    ]
    
    for row_idx, row_data in enumerate(data, start=5):
        for col_idx, value in enumerate(row_data, start=1):
            sheet.cell(row=row_idx, column=col_idx, value=value)
            
    # Add another sheet about conventions
    sheet2 = wb.create_sheet("Conventions")
    sheet2["A1"] = "Procédure"
    sheet2["B1"] = "Description"
    sheet2["A2"] = "Demande de Stage"
    sheet2["B2"] = "L'étudiant doit déposer la demande d'affectation signée par l'entreprise au bureau des stages pour obtenir les conventions. Le traitement prend 48h."
    
    wb.save(os.path.join(DATA_DIR, "procedures_stages.xlsx"))

def main():
    ensure_dir(DATA_DIR)
    create_txt_rules()
    create_docx_inscription()
    create_pdf_calendar()
    create_pdf_bourses()
    create_xlsx_stages()
    print(f"Generated sample documents in {DATA_DIR}")

if __name__ == "__main__":
    main()
